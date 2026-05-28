import json
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches as DocxInches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from app.exporters.file_store import FileStore
from app.schemas.document_ir import DocumentIR
from app.schemas.generation import ExportOption


class OfficeExportService:
    def __init__(self) -> None:
        self.store = FileStore()

    def export_all(self, document: DocumentIR) -> list[ExportOption]:
        exports: list[ExportOption] = []
        if document.type.value in {"presentation", "pitch_deck"}:
            pptx_path = self.store.path_for(document.document_id, document.title, "pptx")
            self.to_pptx(document, pptx_path)
            exports.append(ExportOption(format="pptx", status="ready", url=self.store.url_for(pptx_path)))

        docx_path = self.store.path_for(document.document_id, document.title, "docx")
        self.to_docx(document, docx_path)
        exports.append(ExportOption(format="docx", status="ready", url=self.store.url_for(docx_path)))

        pdf_path = self.store.path_for(document.document_id, document.title, "pdf")
        self.to_pdf(document, pdf_path)
        exports.append(ExportOption(format="pdf", status="ready", url=self.store.url_for(pdf_path)))
        return exports

    def to_pptx(self, document: DocumentIR, path: Path) -> None:
        renderer = Path(__file__).with_name("pptxgen_renderer.mjs")
        payload = document.model_dump(mode="json")
        path.parent.mkdir(parents=True, exist_ok=True)

        with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False, encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False)
            input_path = Path(handle.name)

        try:
            # Find the project root (where package.json and node_modules are)
            # In Docker, it is at /app. Locally it might be 4 levels up.
            project_root = Path(__file__).resolve().parent
            while project_root != project_root.parent:
                if (project_root / "package.json").exists():
                    break
                project_root = project_root.parent

            completed = subprocess.run(
                ["node", str(renderer), str(input_path), str(path.resolve())],
                cwd=str(project_root),
                capture_output=True,
                text=True,
                timeout=90,
                check=False,
            )
            if completed.returncode != 0:
                raise RuntimeError(completed.stderr or completed.stdout or "PptxGenJS renderer failed")
        finally:
            input_path.unlink(missing_ok=True)
    def to_docx(self, document: DocumentIR, path: Path) -> None:
        doc = Document()
        doc.core_properties.title = document.title
        doc.add_heading(document.title, 0)
        doc.add_paragraph(f"Document type: {document.type.value.replace('_', ' ')}")
        doc.add_paragraph(f"Audience: {document.audience.value.replace('_', ' ')}")
        doc.add_paragraph("Generated locally. Editable DOCX export.")

        for section in document.sections:
            doc.add_heading(section.title, level=1)
            for block in section.content_blocks:
                data = block.model_dump()
                if data["type"] == "bullets":
                    for item in data["items"]:
                        doc.add_paragraph(item, style="List Bullet")
                elif data["type"] == "text":
                    doc.add_paragraph(data["text"])
                elif data["type"] == "chart":
                    doc.add_paragraph(f"Chart: {data['title']}")
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Table Grid"
                    table.rows[0].cells[0].text = "Label"
                    table.rows[0].cells[1].text = "Value"
                    for label, value in zip(data["data"]["labels"], data["data"]["values"]):
                        row = table.add_row().cells
                        row[0].text = str(label)
                        row[1].text = str(value)
                elif data["type"] == "diagram":
                    doc.add_paragraph(f"Diagram: {data['title']}")
                    doc.add_paragraph(data["source"])
                elif data["type"] == "table":
                    table = doc.add_table(rows=1, cols=len(data["headers"]))
                    table.style = "Table Grid"
                    for cell, header in zip(table.rows[0].cells, data["headers"]):
                        cell.text = header
                    for row_data in data["rows"]:
                        row = table.add_row().cells
                        for cell, value in zip(row, row_data):
                            cell.text = value
            if section.speaker_notes:
                doc.add_paragraph(f"Notes: {section.speaker_notes}")
        doc.add_page_break()
        doc.add_heading("References", level=1)
        if document.citations:
            for citation in document.citations:
                doc.add_paragraph(citation.formatted_text, style="List Number")
        else:
            doc.add_paragraph("Add verified citations before final submission.")
        doc.save(path)

    def to_pdf(self, document: DocumentIR, path: Path) -> None:
        pdf = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        story = [Paragraph(document.title, styles["Title"]), Spacer(1, 12)]
        story.append(Paragraph(f"Audience: {document.audience.value.replace('_', ' ')}", styles["Normal"]))
        story.append(Spacer(1, 18))
        for section in document.sections:
            story.append(Paragraph(section.title, styles["Heading2"]))
            for block in section.content_blocks:
                data = block.model_dump()
                if data["type"] == "bullets":
                    for item in data["items"]:
                        story.append(Paragraph(f"- {item}", styles["Normal"]))
                elif data["type"] == "text":
                    story.append(Paragraph(data["text"], styles["Normal"]))
                elif data["type"] == "chart":
                    rows = [["Label", "Value"], *[[l, v] for l, v in zip(data["data"]["labels"], data["data"]["values"])] ]
                    table = Table(rows)
                    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)]))
                    story.append(table)
                elif data["type"] == "diagram":
                    story.append(Paragraph(data["title"], styles["Italic"]))
                    story.append(Paragraph(data["source"].replace("\n", "<br/>") , styles["Code"]))
            story.append(Spacer(1, 12))
        pdf.build(story)


